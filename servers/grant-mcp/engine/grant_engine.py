#!/usr/bin/env python3
"""
Grant MCP Python Engine
Handles grant proposal compilation via stdin/stdout JSON protocol.
"""

import json
import sys
import os
from typing import Any


def process_request(request: dict[str, Any]) -> dict[str, Any]:
    """Process a grant compilation request."""
    request_id = request.get("request_id", "unknown")
    source = request.get("source", "")
    sections = request.get("sections", [])
    funder = request.get("funder", "nih")
    program_type = request.get("program_type", "R01")
    bibliography = request.get("bibliography")
    output_format = request.get("output_format", "grant-latex")

    warnings: list[str] = []

    try:
        # For LaTeX output, the TypeScript side handles formatting
        # This engine provides additional compilation support
        if output_format == "docx":
            return generate_docx(sections, funder, program_type, bibliography, warnings)
        else:
            return {
                "request_id": request_id,
                "status": "success",
                "message": "LaTeX generation handled by TypeScript formatter",
                "warnings": warnings,
            }
    except Exception as e:
        return {
            "request_id": request_id,
            "status": "error",
            "error": {
                "code": "GRANT_ENGINE_ERROR",
                "message": str(e),
            },
        }


def generate_docx(
    sections: list[dict[str, Any]],
    funder: str,
    program_type: str,
    bibliography: str | None,
    warnings: list[str],
) -> dict[str, Any]:
    """Generate DOCX output using python-docx."""
    try:
        from docx import Document  # type: ignore[import-untyped]
        from docx.shared import Pt, Inches  # type: ignore[import-untyped]
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
        import base64
        import tempfile

        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.size = Pt(11)
        if funder == "nih":
            font.name = "Arial"
        else:
            font.name = "Times New Roman"

        # Set margins
        for section in doc.sections:
            if funder == "nih":
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            else:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

        # Add sections
        for sec in sections:
            role = sec.get("role", "other")
            content = sec.get("content", "")
            title = sec.get("title", role.replace("-", " ").title())

            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

            for para_text in content.split("\n\n"):
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())

        # Save to temporary file and read as base64
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            with open(tmp.name, "rb") as f:
                docx_base64 = base64.b64encode(f.read()).decode("utf-8")
            os.unlink(tmp.name)

        return {
            "status": "success",
            "output_base64": docx_base64,
            "format": "docx",
            "warnings": warnings,
        }

    except ImportError:
        warnings.append("python-docx not installed. Install with: pip install python-docx")
        return {
            "status": "error",
            "error": {
                "code": "MISSING_DEPENDENCY",
                "message": "python-docx is required for DOCX output",
            },
            "warnings": warnings,
        }


def main() -> None:
    """Main entry point: read JSON from stdin, write JSON to stdout."""
    try:
        raw_input = sys.stdin.read()
        if not raw_input.strip():
            result = {
                "status": "error",
                "error": {"code": "EMPTY_INPUT", "message": "No input provided"},
            }
        else:
            request = json.loads(raw_input)
            result = process_request(request)
    except json.JSONDecodeError as e:
        result = {
            "status": "error",
            "error": {"code": "INVALID_JSON", "message": str(e)},
        }
    except Exception as e:
        result = {
            "status": "error",
            "error": {"code": "UNEXPECTED_ERROR", "message": str(e)},
        }

    json.dump(result, sys.stdout, indent=2)
    sys.stdout.flush()


if __name__ == "__main__":
    main()
